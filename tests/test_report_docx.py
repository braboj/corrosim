"""The Word (.docx) report renderer: builds a valid, re-openable document that
mirrors the HTML report's data, equations and validation content."""
from __future__ import annotations

import matplotlib

matplotlib.use("Agg")

import pandas as pd
import pytest

pytest.importorskip("docx")            # python-docx is the `report`/`dev` extra
from docx import Document
from docx.oxml.ns import qn

from corrosim import equations, report, report_docx
from corrosim.speciation import analyse_speciation, protonation_fraction


def _native_equation_count(doc) -> int:
    """Number of native (editable) OMML equations in the document body."""
    return sum(len(p._p.findall(qn("m:oMath"))) for p in doc.paragraphs)


def _row(name, gap, hardness):
    return {
        "name": name, "formula": "C15H10O6", "charge": 0,
        "level": "B3LYP/6-311++G(d,p) (ddCOSMO:water)",
        "homo_ev": -6.0, "lumo_ev": -6.0 + gap, "gap_ev": gap,
        "hardness_ev": hardness, "softness_inv_ev": 1 / hardness,
        "electronegativity_ev": 4.0, "electrophilicity_ev": 4.0,
        "delta_n": 0.2, "back_donation_ev": -0.5, "tnc": -4.0,
    }


def _all_text(doc) -> str:
    paras = "\n".join(p.text for p in doc.paragraphs)
    cells = "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    return paras + "\n" + cells


def test_build_docx_report_is_valid_and_mirrors_content(tmp_path):
    neutral = [_row("quercetin", 4.0, 2.0), _row("kaempferol", 4.4, 2.2)]
    mc = [{"name": "quercetin", "e_ads_kjmol": -16.0},
          {"name": "kaempferol", "e_ads_kjmol": -16.5}]
    md = [{"name": "quercetin", "metal": "Fe", "metal_O_peak_A": 3.65},
          {"name": "kaempferol", "metal": "Fe", "metal_O_peak_A": 3.35}]
    fukui = {"quercetin": [{"idx": 5, "symbol": "O", "f_minus": 0.09, "f_plus": 0.1}]}
    protonated = [{**_row("quercetin+H+", 3.3, 1.6), "delta_n": -0.05},
                  {**_row("kaempferol+H+", 3.6, 1.8), "delta_n": -0.07}]
    summary = analyse_speciation(
        neutral, protonated, ph=0.0,
        rank_fn=lambda r: report.rank_inhibitors(pd.DataFrame(r)).to_dict("records"))
    computed = [{"name": n, "pkah": pk, "f_protonated": protonation_fraction(0.0, pk)}
                for n, pk in [("quercetin", -12.1), ("kaempferol", -11.2)]]

    out = tmp_path / "report.docx"
    res = report_docx.build_docx_report(
        neutral, mc, md, fukui, figdir=str(tmp_path / "nope"), out_path=str(out),
        medium="1 M HCl", order=["quercetin", "kaempferol"],
        acid_cation_rows=protonated, speciation_summary=summary,
        computed_pkah=computed, opt_neutral_rows=neutral, opt_acid_rows=protonated,
        generated_at="2026-01-01 00:00")
    assert res == str(out)
    assert out.exists() and out.stat().st_size > 1000

    doc = Document(str(out))             # re-opens => structurally valid
    text = _all_text(doc)
    # the molecules, the sections, and the woven validation numbers are present
    assert "quercetin" in text and "kaempferol" in text
    assert "Scientific basis & validation" in text
    assert "Speciation in 1 M HCl" in text
    assert "99.62" in text               # experimental anchor (Mohammed 2014)
    assert "-16.0" in text               # Stage-2 adsorption merged in
    # equations are native, editable Word equations (OMML), not images
    assert _native_equation_count(doc) > 10


def test_build_docx_report_without_optional_sections(tmp_path):
    # minimal inputs (no acid/speciation/opt) must still produce a valid file
    neutral = [_row("quercetin", 4.0, 2.0)]
    out = tmp_path / "min.docx"
    report_docx.build_docx_report(neutral, [], [], {}, figdir=str(tmp_path / "x"),
                                  out_path=str(out), order=["quercetin"])
    doc = Document(str(out))
    assert "quercetin" in _all_text(doc)


def test_latex_to_omml_converts_every_equation():
    pytest.importorskip("latex2mathml")
    pytest.importorskip("mathml2omml")
    for key, eq in equations.EQUATIONS.items():
        el = report_docx._latex_to_omml(eq.latex)
        assert el is not None, f"{key} failed to convert to OMML"
        assert el.tag.endswith("}oMath")


def test_equation_falls_back_to_image_without_toolchain(tmp_path, monkeypatch):
    # if the LaTeX->OMML toolchain is unavailable, the equation degrades to the
    # rendered image rather than vanishing.
    monkeypatch.setattr(report_docx, "_latex_to_omml", lambda _latex: None)
    d = report_docx._Doc()
    d.equation("delta_n")
    assert len(d.doc.inline_shapes) == 1                 # image fallback present
    assert _native_equation_count(d.doc) == 0
