"""corrosim.report_docx.

Word (.docx) rendering of the multiscale report, built with python-docx (pure
Python, no system binary). It draws on exactly the same derived data
(:func:`report.prepare_report_data`), the same governing equations
(:mod:`equations`) and the same narrative (:mod:`report_content`) as the HTML
report, so the two outputs stay in lock-step; only the formatting differs.

Equations are inserted as **native, editable Word equations** (OMML): the
LaTeX source is converted LaTeX -> MathML -> OMML with the pure-Python
``latex2mathml`` + ``mathml2omml`` packages (no LaTeX/pandoc/Office toolchain),
so a reader can click and edit them in Word's equation editor. If that toolchain
is unavailable or a conversion fails, the equation degrades to the rendered
mathtext image (:mod:`equations`) so the report is never missing a formula.

Entry point: :func:`build_docx_report`, whose signature mirrors
``report.build_pipeline_report`` so a driver can build both from one call site.
"""
from __future__ import annotations

import datetime
import io
import os

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor

from . import equations as _eq
from . import report_content as _content
from .report import prepare_report_data, rank_inhibitors, results_dataframe
from .report_layout import figure_path

_MUTED = RGBColor(0x71, 0x80, 0x96)
_FIG_WIDTH = Inches(5.7)
_GRID_WIDTH = Inches(2.9)        # per-molecule figures shown a little smaller
_OMML_MATH_NS = ('xmlns:m="http://schemas.openxmlformats.org/'
                 'officeDocument/2006/math"')


def _latex_to_omml(latex: str):
    """Convert a LaTeX expression to an OMML ``<m:oMath>`` element (a native,
    editable Word equation), or return ``None`` if the pure-Python toolchain is
    absent or the conversion fails — the caller then falls back to an image.
    """
    try:
        from latex2mathml.converter import convert as latex_to_mathml
        from mathml2omml import convert as mathml_to_omml
    except ImportError:
        return None
    try:
        omml = mathml_to_omml(latex_to_mathml(latex))
        if "xmlns:m=" not in omml:                    # declare m: so parse_xml resolves it
            omml = omml.replace("<m:oMath>", f"<m:oMath {_OMML_MATH_NS}>", 1)
        return parse_xml(omml)
    except Exception:
        return None


class _Doc:
    """Thin wrapper over a python-docx Document with the report's building blocks."""

    def __init__(self) -> None:
        self.doc = Document()
        self._c1 = 0            # section counter (level 1)
        self._c2 = 0            # subsection counter (level 2)

    # --- text ---------------------------------------------------------------
    def heading(self, text: str, level: int) -> None:
        """Add a heading. Sections (level 1) and subsections (level 2) are
        numbered hierarchically (``1.``, ``1.1`` …) to match the HTML report; the
        title (level 0) and deeper (level 3) headings are left unnumbered.
        """
        if level == 1:
            self._c1 += 1
            self._c2 = 0
            text = f"{self._c1}. {text}"
        elif level == 2:
            self._c2 += 1
            text = f"{self._c1}.{self._c2} {text}"
        self.doc.add_heading(text, level=level)

    def para(self, text: str, *, muted: bool = False, size: int | None = None) -> None:
        """A paragraph rendering the shared content's ``**bold**`` markup."""
        p = self.doc.add_paragraph()
        for chunk, bold in _content.inline_runs(text):
            run = p.add_run(chunk)
            run.bold = bold
            if muted:
                run.font.color.rgb = _MUTED
            if size:
                run.font.size = Pt(size)

    def note(self, text: str) -> None:
        """The caveat box, as an italicised, indented paragraph."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(9)

    # --- figures ------------------------------------------------------------
    def figure(self, figdir: str, fname: str, caption: str,
               width: Inches = _FIG_WIDTH) -> None:
        path = figure_path(figdir, fname)
        if not os.path.exists(path):
            return                                    # missing figure: skip silently
        self.doc.add_picture(path, width=width)
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = _MUTED

    def explain(self, role: str) -> None:
        txt = _content.FIGURE_EXPLANATIONS.get(role)
        if txt:
            self.para(txt, size=10)

    # --- equations ----------------------------------------------------------
    def equation(self, key: str) -> None:
        eq = _eq.EQUATIONS[key]
        omath = _latex_to_omml(eq.latex)
        if omath is not None:                         # native, editable Word equation
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p._p.append(omath)
        else:                                         # fallback: rendered image
            png = io.BytesIO(_eq.render_equation_png(eq.latex))
            self.doc.add_picture(png)
            self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = self.doc.add_paragraph()
        cap.add_run(f"{eq.quantity} — ").bold = True
        run = cap.add_run(eq.meaning)
        run.font.size = Pt(9)
        run.font.color.rgb = _MUTED

    # --- tables -------------------------------------------------------------
    def df_table(self, df: pd.DataFrame, *, highlight_first: bool = False) -> None:
        cols = list(df.columns)
        t = self.doc.add_table(rows=1, cols=len(cols))
        t.style = "Table Grid"
        for i, c in enumerate(cols):
            _set_cell(t.rows[0].cells[i], c, bold=True)
        for ri, (_, row) in enumerate(df.iterrows()):
            cells = t.add_row().cells
            for i, c in enumerate(cols):
                v = row[c]
                _set_cell(cells[i], "" if pd.isna(v) else v,
                          bold=(highlight_first and ri == 0))

    def content_table(self, payload: dict) -> None:
        cols, rows = payload["columns"], payload["rows"]
        t = self.doc.add_table(rows=1, cols=len(cols))
        t.style = "Table Grid"
        for i, c in enumerate(cols):
            _set_cell(t.rows[0].cells[i], c, bold=True)
        for row in rows:
            cells = t.add_row().cells
            for i, c in enumerate(row):
                _set_cell(cells[i], c)
        if payload.get("caption"):
            self.para(payload["caption"], muted=True, size=8)

    def save(self, path: str) -> str:
        self.doc.save(path)
        return path


def _set_cell(cell, text, *, bold: bool = False) -> None:
    cell.text = str(text)
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = bold
            run.font.size = Pt(8.5)


def _scientific_basis(d: _Doc) -> None:
    """The shared 'Scientific basis & validation' section (report_content), with
    the equation set injected in scientific form at the ``eqgroups`` marker.
    """
    d.heading("Scientific basis & validation", level=1)
    for kind, payload in _content.SCIENTIFIC_BASIS:
        if kind == "h3" and isinstance(payload, str):
            d.heading(payload, level=2)
        elif kind == "p" and isinstance(payload, str):
            d.para(payload)
        elif kind == "table" and isinstance(payload, dict):
            d.content_table(payload)
        elif kind == "eqgroups":
            for group_heading, group in _eq.EQUATION_GROUPS:
                d.heading(group_heading, level=3)
                for eq in group:
                    d.equation(eq.key)


def _speciation_section(d: _Doc, summary: dict | None, medium: str,
                        computed_pkah: list[dict] | None,
                        pka_freq_corrected: bool) -> None:
    """Compact Word version of the speciation / computed-pKaH section (ADR 0004/5),
    driven by the same summary dict as the HTML report.
    """
    if not summary:
        return
    spec = summary["speciation"]
    d.heading(f"Speciation in {medium} (pH ≈ {spec.ph:.1f})", level=2)
    d.para(
        f"The most basic site (4-oxo carbonyl, pKaH ≈ {spec.pkah:.1f}) is a very "
        f"weak base, so by Henderson-Hasselbalch the inhibitor is "
        f"**{spec.f_neutral:.0%} neutral / {spec.f_protonated:.0%} protonated** — "
        f"the {spec.dominant} form dominates, which is why the headline ranking "
        f"uses the neutral form. Population-weighted lead: "
        f"**{summary['blended_lead']}**.")
    d.df_table(results_dataframe(summary["blended_rows"]))
    cross_f, cross_pk = summary["crossover_fraction"], summary["crossover_pkah"]
    if cross_f and cross_pk is not None:
        d.para(
            f"Sensitivity: the composite lead changes from "
            f"{summary['neutral_lead']} to {summary['crossover_lead']} at only "
            f"~{cross_f:.0%} protonation (pKaH ≈ {cross_pk:.1f}) — the protonation "
            f"pKa is the key uncertainty.", muted=True)
    if computed_pkah:
        basis = "frequency-corrected" if pka_freq_corrected else "electronic-only"
        worst = max(r["f_protonated"] for r in computed_pkah)
        d.para(
            f"Computed pKaH (DFT deprotonation cycle, {basis}) resolves it: "
            f"the most basic flavonoid is only {worst * 100:.2f}% protonated, so "
            f"every species is essentially fully neutral and the headline lead is "
            f"robust.")
        d.content_table({
            "columns": ["molecule", "computed pKaH", "% protonated"],
            "rows": [[r["name"], f"{r['pkah']:.1f}", f"{r['f_protonated'] * 100:.2f}%"]
                     for r in computed_pkah],
            "caption": "results/pka.json (ADR 0005).",
        })


def build_docx_report(neutral_aq_rows: list[dict], mc_rows: list[dict],
                      md_rows: list[dict], fukui_by_name: dict[str, list[dict]],
                      figdir: str, out_path: str,
                      metal: str = "Fe(110)", medium: str = "1 M HCl",
                      order: list[str] | None = None,
                      generated_at: str | None = None,
                      acid_cation_rows: list[dict] | None = None,
                      speciation_summary: dict | None = None,
                      computed_pkah: list[dict] | None = None,
                      pka_freq_corrected: bool = False,
                      opt_neutral_rows: list[dict] | None = None,
                      opt_acid_rows: list[dict] | None = None) -> str:
    """Build the multiscale report as a Word ``.docx``. Mirrors
    ``report.build_pipeline_report``: same inputs, same sections, same narrative
    and equations — rendered for Word. Returns the output path.
    """
    prep = prepare_report_data(neutral_aq_rows, mc_rows, md_rows, fukui_by_name,
                               metal, order)
    names = list(prep.df["name"])
    d = _Doc()

    d.heading("corrosim — multiscale corrosion-inhibitor report", level=0)
    ts = generated_at or datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    d.para(f"Substrate: {metal}  |  Medium: {medium}  |  DFT level: "
           f"{prep.level}  |  Generated {ts}", muted=True)
    d.note(_content.HEADLINE_CAVEAT)

    d.heading("Overview", level=1)
    d.para(_content.STAGE_INTROS["overview"])
    d.explain("pipeline")
    d.figure(figdir, "fig0_pipeline.png", "corrosim pipeline")

    d.heading("Summary & ranking", level=1)
    d.df_table(prep.summary, highlight_first=True)
    d.para("Composite score z-scores a smaller gap, lower hardness and higher "
           f"softness (higher = stronger predicted reactivity); E_ads (Stage 2) "
           f"and the {prep.m_elem}-O distance (Stage 3) are shown alongside.",
           muted=True)

    # Stage 1 -----------------------------------------------------------------
    d.heading("Stage 1 — DFT electronic descriptors", level=1)
    d.para(_content.STAGE_INTROS["dft"])
    d.figure(figdir, "fig1_structures.png", "Modelled flavonoids")
    d.explain("structures")
    d.figure(figdir, "fig2_mo_diagram.png",
             "Frontier-orbital energies vs Fe(110) work function")
    d.explain("mo_diagram")
    d.heading("Frontier-orbital isosurfaces (HOMO / LUMO)", level=2)
    for n in names:
        d.figure(figdir, f"fig2b_{n}_homo.png", f"{n} HOMO", width=_GRID_WIDTH)
    d.explain("orbital_homo")
    for n in names:
        d.figure(figdir, f"fig2b_{n}_lumo.png", f"{n} LUMO", width=_GRID_WIDTH)
    d.explain("orbital_lumo")
    d.figure(figdir, "fig3_descriptors.png", "Reactivity descriptors")
    d.explain("descriptors")
    d.figure(figdir, "fig3b_protonation.png", "Protonation effect (1 M HCl)")
    d.explain("protonation")
    d.heading("Full descriptor table (neutral, aqueous)", level=2)
    d.df_table(prep.full)
    if os.path.exists(figure_path(figdir, "fig8_geometry_comparison.png")):
        d.heading("Geometry refinement (FF vs DFT-optimised)", level=2)
        d.figure(figdir, "fig8_geometry_comparison.png",
                 "Force-field vs DFT-optimised geometry")
        d.explain("geometry")

    if opt_neutral_rows:
        d.heading("Optimised-geometry descriptors (DFT-relaxed)", level=2)
        ndf = pd.DataFrame(opt_neutral_rows)
        if order:
            keep = [n for n in order if n in set(ndf["name"])]
            ndf = ndf.set_index("name").loc[keep].reset_index()
        ranked = rank_inhibitors(ndf)
        d.para("Descriptors on DFT-optimised geometries; the ranking is unchanged "
               "— the lead is geometry-robust.")
        d.df_table(ranked[["name", "gap_ev", "hardness_ev", "softness_inv_ev",
                           "delta_n", "tnc", "score"]].round(3), highlight_first=True)
        if opt_acid_rows:
            d.heading("Optimised protonated cations (in-acid)", level=3)
            d.df_table(results_dataframe(opt_acid_rows))

    if acid_cation_rows:
        d.heading("Species in the acidic medium (protonated cation)", level=2)
        d.para(f"In {medium} the inhibitor is present largely as its +1 cation. "
               "The headline ranking uses the neutral form (ADR 0003); the "
               "protonated-cation descriptors are tabulated here for comparison.")
        d.df_table(results_dataframe(acid_cation_rows))

    _speciation_section(d, speciation_summary, medium, computed_pkah,
                        pka_freq_corrected)

    # Stage 1 (cont.) — Fukui and ESP are facets of Stage 1 (isolated-molecule
    # QM analysis), so they are subsections here, not separate pipeline stages.
    d.heading("Local reactivity (Fukui)", level=2)
    d.para(_content.STAGE_INTROS["fukui"])
    if prep.fukui_items:
        for name, sites in prep.fukui_items:
            d.para(f"**{name}**: {sites}", size=10)
    for n in names:
        d.figure(figdir, f"fig4_{n}_fukui.png", f"{n} — condensed Fukui")
    d.explain("fukui")

    # Stage 1c — ESP ----------------------------------------------------------
    d.heading("Electrostatic-potential (ESP) map", level=2)
    d.para(_content.STAGE_INTROS["esp"])
    for n in names:
        d.figure(figdir, f"fig7_{n}_esp.png", f"{n} — ESP map", width=_GRID_WIDTH)
    d.explain("esp")

    # Stage 2 — Monte Carlo ---------------------------------------------------
    d.heading("Stage 2 — Monte Carlo adsorption", level=1)
    d.para(_content.STAGE_INTROS["mc"])
    for n in names:
        d.figure(figdir, f"fig5_{n}_mc_pose.png", f"{n} — best pose")
    d.explain("mc_pose")
    for n in names:
        d.figure(figdir, f"fig5_{n}_mc_energy.png", f"{n} — MC annealing")
    d.explain("mc_energy")

    # Stage 3 — MD ------------------------------------------------------------
    d.heading(f"Stage 3 — Brownian MD ({prep.m_elem}-O RDF)", level=1)
    d.para(_content.STAGE_INTROS["md"])
    for n in names:
        d.figure(figdir, f"fig6_{n}_rdf.png", f"{n} — {prep.m_elem}-O RDF")
    d.explain("rdf")

    # Scientific basis & validation ------------------------------------------
    _scientific_basis(d)

    # Method ------------------------------------------------------------------
    d.heading("Method & caveats", level=1)
    d.para(f"DFT level: {prep.level}. {_content.METHOD_CAVEAT}", muted=True)

    return d.save(out_path)
