"""corrosim.report_docx.

Word (.docx) rendering of the multiscale report, built with python-docx (pure
Python, no system binary). It draws on exactly the same derived data
(:func:`report.prepare_report_data`) and the same shared caveats / bottom-line
(:mod:`report_content`) as the HTML report, so the two outputs stay in
lock-step; only the formatting differs.

The report is lean: tables and figures under each stage with minimal captions,
no methodology essay. The full methodology lives in ``docs/pipeline.md`` and the
validation record in ``docs/validation.md``.

Entry point: :func:`build_docx_report`, whose signature mirrors
``report.build_pipeline_report`` so a driver can build both from one call site.
"""
from __future__ import annotations

import datetime
import os
from typing import TYPE_CHECKING

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from . import report_content as _content
from .report import prepare_report_data, rank_inhibitors, results_dataframe
from .report_layout import figure_path

if TYPE_CHECKING:
    from .report import PreparedReport

_MUTED = RGBColor(0x71, 0x80, 0x96)
_FIG_WIDTH = Inches(5.7)
# Per-molecule figures shown a little smaller
_GRID_WIDTH = Inches(2.9)


class _Doc:
    """Thin wrapper over a python-docx Document with the report's blocks."""

    def __init__(self) -> None:
        self.doc = Document()
        # Section counter (level 1)
        self._c1 = 0
        # Subsection counter (level 2)
        self._c2 = 0

    # --- text ---------------------------------------------------------------
    def heading(self, text: str, level: int) -> None:
        """Add a heading. Sections (level 1) and subsections (level 2) are
        numbered hierarchically (``1.``, ``1.1`` …) to match the HTML report;
        the title (level 0) and deeper (level 3) headings are left unnumbered.
        """
        if level == 1:
            self._c1 += 1
            self._c2 = 0
            text = f"{self._c1}. {text}"
        elif level == 2:
            self._c2 += 1
            text = f"{self._c1}.{self._c2} {text}"
        self.doc.add_heading(text, level=level)

    def para(self, text: str, *, muted: bool = False,
             size: int | None = None) -> None:
        """A paragraph rendering the shared content's ``**bold**`` markup."""
        p = self.doc.add_paragraph()
        for chunk, bold in _content.inline_runs(text):
            run = p.add_run(chunk)
            run.bold = bold
            if muted:
                run.font.color.rgb = _MUTED
            if size:
                run.font.size = Pt(size)

    def note(self, text: str) -> None:
        """The caveat box, as an italicised, indented paragraph."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(9)

    # --- figures ------------------------------------------------------------
    def figure(self, figdir: str, fname: str, caption: str,
               width: Inches = _FIG_WIDTH) -> None:
        path = figure_path(figdir, fname)
        if not os.path.exists(path):
            # Missing figure: skip silently
            return
        self.doc.add_picture(path, width=width)
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = _MUTED

    # --- tables -------------------------------------------------------------
    def df_table(self, df: pd.DataFrame, *,
                 highlight_first: bool = False) -> None:
        cols = list(df.columns)
        t = self.doc.add_table(rows=1, cols=len(cols))
        t.style = "Table Grid"
        for i, c in enumerate(cols):
            _set_cell(t.rows[0].cells[i], c, bold=True)
        for ri, (_, row) in enumerate(df.iterrows()):
            cells = t.add_row().cells
            for i, c in enumerate(cols):
                v = row[c]
                _set_cell(cells[i], "" if pd.isna(v) else v,
                          bold=(highlight_first and ri == 0))

    def content_table(self, payload: dict) -> None:
        cols, rows = payload["columns"], payload["rows"]
        t = self.doc.add_table(rows=1, cols=len(cols))
        t.style = "Table Grid"
        for i, c in enumerate(cols):
            _set_cell(t.rows[0].cells[i], c, bold=True)
        for row in rows:
            cells = t.add_row().cells
            for i, c in enumerate(row):
                _set_cell(cells[i], c)
        if payload.get("caption"):
            self.para(payload["caption"], muted=True, size=8)

    def save(self, path: str) -> str:
        self.doc.save(path)
        return path


def _set_cell(cell, text, *, bold: bool = False) -> None:
    cell.text = str(text)
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = bold
            run.font.size = Pt(8.5)


def _speciation_section(d: _Doc, summary: dict | None, medium: str,
                        computed_pkah: list[dict] | None,
                        pka_freq_corrected: bool) -> None:
    """Compact Word speciation + computed-pKaH tables, from the same summary
    dict as the HTML report.
    """
    if not summary:
        return
    spec = summary["speciation"]
    d.heading(f"Speciation in {medium} (pH ≈ {spec.ph:.1f})", level=2)
    d.para(
        f"**{spec.f_neutral:.0%} neutral / {spec.f_protonated:.0%} "
        f"protonated** at this pH — the {spec.dominant} form dominates. "
        f"Population-weighted descriptors (blended lead "
        f"**{summary['blended_lead']}**):", muted=True, size=9)
    d.df_table(results_dataframe(summary["blended_rows"]))
    if computed_pkah:
        basis = ("frequency-corrected" if pka_freq_corrected
                 else "electronic-only")
        d.content_table({
            "columns": ["molecule", "computed pKaH", "% protonated"],
            "rows": [[r["name"], f"{r['pkah']:.1f}",
                      f"{r['f_protonated'] * 100:.2f}%"]
                     for r in computed_pkah],
            "caption": f"Computed pKaH (DFT cycle, {basis}).",
        })


def _docx_header(d: _Doc, prep: PreparedReport, metal: str, medium: str,
                 generated_at: str | None) -> None:
    """Title, run-metadata line and the headline caveat.

    The data-derived headline sentence is placed in the Summary & ranking
    section (:func:`_summary_section`), not here.
    """
    d.heading("corrosim — multiscale corrosion-inhibitor report", level=0)
    ts = generated_at or datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    d.para(f"Substrate: {metal}  |  Medium: {medium}  |  DFT level: "
           f"{prep.level}  |  Generated {ts}", muted=True)
    d.note(_content.HEADLINE_CAVEAT)


def _overview_section(d: _Doc, figdir: str) -> None:
    """Overview heading + the pipeline diagram (methodology in pipeline.md)."""
    d.heading("Overview", level=1)
    d.figure(figdir, "fig0_pipeline.png", "corrosim pipeline")


def _summary_section(d: _Doc, prep: PreparedReport) -> None:
    """Headline sentence + the ranking table + the one-line scoring note."""
    d.heading("Summary & ranking", level=1)
    bl = prep.bottom_line()
    if bl:
        d.para(bl)
    d.df_table(prep.summary, highlight_first=True)
    d.para(_content.score_note(prep.m_elem), muted=True, size=9)


def _dft_section(d: _Doc, prep: PreparedReport, figdir: str,
                 names: list[str]) -> None:
    """DFT descriptors: structures, MO diagram, per-molecule HOMO/LUMO
    isosurfaces, descriptor charts, the full table and the optional geometry-
    refinement figure.
    """
    d.heading("DFT electronic descriptors", level=1)
    d.figure(figdir, "fig1_structures.png", "Modelled molecules")
    d.figure(figdir, "fig2_mo_diagram.png",
             "Frontier-orbital energies vs the metal work function")
    d.heading("Frontier-orbital isosurfaces (HOMO / LUMO)", level=2)
    for n in names:
        d.figure(figdir, f"fig2b_{n}_homo.png", f"{n} HOMO", width=_GRID_WIDTH)
    for n in names:
        d.figure(figdir, f"fig2b_{n}_lumo.png", f"{n} LUMO", width=_GRID_WIDTH)
    d.figure(figdir, "fig3_descriptors.png", "Reactivity descriptors")
    d.figure(figdir, "fig3b_protonation.png", "Protonation effect")
    d.heading("Full descriptor table (neutral, aqueous)", level=2)
    d.df_table(prep.full)
    if os.path.exists(figure_path(figdir, "fig8_geometry_comparison.png")):
        d.heading("Geometry refinement (FF vs DFT-optimised)", level=2)
        d.figure(figdir, "fig8_geometry_comparison.png",
                 "Force-field vs DFT-optimised geometry")


def _optimised_geometry_section(d: _Doc, opt_neutral_rows: list[dict] | None,
                                opt_acid_rows: list[dict] | None,
                                order: list[str] | None) -> None:
    """Optional DFT-optimised-geometry descriptor tables (neutral + cation)."""
    if not opt_neutral_rows:
        return
    d.heading("Optimised-geometry descriptors (DFT-relaxed)", level=2)
    ndf = pd.DataFrame(opt_neutral_rows)
    if order:
        keep = [n for n in order if n in set(ndf["name"])]
        ndf = ndf.set_index("name").loc[keep].reset_index()
    ranked = rank_inhibitors(ndf)
    d.df_table(ranked[["name", "gap_ev", "hardness_ev", "softness_inv_ev",
                       "delta_n", "tnc", "score"]].round(3),
               highlight_first=True)
    if opt_acid_rows:
        d.heading("Optimised protonated cations (in-acid)", level=3)
        d.df_table(results_dataframe(opt_acid_rows))


def _acid_cation_section(d: _Doc, acid_cation_rows: list[dict] | None,
                         medium: str) -> None:
    """Optional protonated-cation descriptor table for the acidic medium."""
    if not acid_cation_rows:
        return
    d.heading("Species in the acidic medium (protonated cation)", level=2)
    d.df_table(results_dataframe(acid_cation_rows))
    d.para(f"Protonated +1 cation descriptors in {medium}; the headline "
           "ranking stays on the neutral form (see docs/pipeline.md).",
           muted=True, size=9)


def _fukui_section(d: _Doc, prep: PreparedReport, figdir: str,
                   names: list[str]) -> None:
    """Local-reactivity (Fukui) subsection: donor sites + per-molecule maps."""
    d.heading("Local reactivity (Fukui)", level=2)
    if prep.fukui_items:
        for name, sites in prep.fukui_items:
            d.para(f"**{name}**: {sites}", size=10)
    for n in names:
        d.figure(figdir, f"fig4_{n}_fukui.png", f"{n} — condensed Fukui")


def _esp_section(d: _Doc, figdir: str, names: list[str]) -> None:
    """Electrostatic-potential (ESP) map subsection."""
    d.heading("Electrostatic-potential (ESP) map", level=2)
    for n in names:
        d.figure(figdir, f"fig7_{n}_esp.png", f"{n} — ESP map",
                 width=_GRID_WIDTH)


def _mc_section(d: _Doc, figdir: str, names: list[str]) -> None:
    """Monte Carlo adsorption: per-molecule pose + annealing figures."""
    d.heading("Monte Carlo adsorption", level=1)
    for n in names:
        d.figure(figdir, f"fig5_{n}_mc_pose.png", f"{n} — best pose")
    for n in names:
        d.figure(figdir, f"fig5_{n}_mc_energy.png", f"{n} — MC annealing")


def _md_section(d: _Doc, prep: PreparedReport, figdir: str,
                names: list[str]) -> None:
    """Brownian-MD metal-O RDF subsection."""
    d.heading(f"Brownian MD ({prep.m_elem}-O RDF)", level=1)
    for n in names:
        d.figure(figdir, f"fig6_{n}_rdf.png", f"{n} — {prep.m_elem}-O RDF")


def _method_section(d: _Doc, prep: PreparedReport) -> None:
    """Method & caveats footer."""
    d.heading("Method & caveats", level=1)
    d.para(f"DFT level: {prep.level}. {_content.METHOD_CAVEAT}", muted=True)


# Orchestrates the Word report section by section, delegating each to a
# _*_section builder (mirrors the HTML report's section-by-section build).
def build_docx_report(
        neutral_aq_rows: list[dict], mc_rows: list[dict],
        md_rows: list[dict], fukui_by_name: dict[str, list[dict]],
        figdir: str, out_path: str,
        metal: str = "Fe(110)", medium: str = "1 M HCl",
        order: list[str] | None = None,
        generated_at: str | None = None,
        acid_cation_rows: list[dict] | None = None,
        speciation_summary: dict | None = None,
        computed_pkah: list[dict] | None = None,
        pka_freq_corrected: bool = False,
        opt_neutral_rows: list[dict] | None = None,
        opt_acid_rows: list[dict] | None = None) -> str:
    """Build the multiscale report as a Word ``.docx``.

    Mirrors ``report.build_pipeline_report``: same inputs, same lean sections
    (tables + figures under each stage) — rendered for Word.

    Args:
        neutral_aq_rows: Neutral aqueous descriptor rows.
        mc_rows: Monte Carlo adsorption summary rows.
        md_rows: Brownian-MD RDF summary rows.
        fukui_by_name: Per-molecule Fukui rows keyed by name.
        figdir: The figures root directory.
        out_path: Destination ``.docx`` path.
        metal: Substrate label.
        medium: Corrosive medium label.
        order: Molecule display order (defaults to the input order).
        generated_at: Timestamp string (defaults to now).
        acid_cation_rows: Protonated-cation descriptor rows, if any.
        speciation_summary: The pH-speciation summary dict, if any.
        computed_pkah: Computed-pKaH rows, if any.
        pka_freq_corrected: Whether the pKaH is frequency-corrected.
        opt_neutral_rows: DFT-optimised neutral rows, if any.
        opt_acid_rows: DFT-optimised protonated-cation rows, if any.

    Returns:
        The output ``.docx`` path.
    """
    prep = prepare_report_data(neutral_aq_rows, mc_rows, md_rows, fukui_by_name,
                               metal, order)
    names = list(prep.df["name"])
    d = _Doc()
    _docx_header(d, prep, metal, medium, generated_at)
    _overview_section(d, figdir)
    _summary_section(d, prep)
    _dft_section(d, prep, figdir, names)
    _optimised_geometry_section(d, opt_neutral_rows, opt_acid_rows, order)
    _acid_cation_section(d, acid_cation_rows, medium)
    _speciation_section(d, speciation_summary, medium, computed_pkah,
                        pka_freq_corrected)
    # Fukui and ESP are facets of the isolated-molecule QM analysis, so they
    # render as subsections here, not as separate pipeline stages.
    _fukui_section(d, prep, figdir, names)
    _esp_section(d, figdir, names)
    _mc_section(d, figdir, names)
    _md_section(d, prep, figdir, names)
    _method_section(d, prep)
    return d.save(out_path)
